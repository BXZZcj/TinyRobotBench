import sys
from config.prompt_config import *
from config.model_config import *
from dao.knowledge_dao import KnowledgeDAO
from util.llm_request import request_llm
from util.embedding import calculate_embedding
import faiss
import chainlit as cl
import asyncio

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.embeddings.sentence_transformer import SentenceTransformerEmbeddings
from langchain import hub
from langchain_core.documents.base import Document
import docx
import io
from PIL import Image
import base64


# @singleton
class ChainlitReasoner:
    def __init__(self):
        self.knowledge_dao = KnowledgeDAO()
        self.qa_knowledge = self.knowledge_dao.load_qa_knowledge()
        self.faiss_index = faiss.IndexFlatL2(self.qa_knowledge['question_embeddings'].shape[1])
        self.faiss_index.add(self.qa_knowledge['question_embeddings'])
        self.response_history = []

    def test(self):
        ...

    async def simple_answer(self, user_question, response_history, **gen_kwargs):
        '''这里传给LLM的prompt是光秃秃的user_question, 其实理应加一些prompt template的'''
        llm_response = request_llm(user_question.content, response_history, **gen_kwargs)
        await cl.Message(content=llm_response['response'].replace('</s>', '')).send()
        return llm_response['response'], llm_response['history']

    async def reason_answer(self, user_question, response_history, knowledge_top_k=3, **gen_kwargs):
        root_id = await cl.Message(content="Reasoning from knowledge...").send()
        await cl.Message(author="Question", content=f"{user_question}", parent_id=root_id).send()

        await cl.Message(author="Faiss", content=f'Searching related questions by calculating vector similarity.', parent_id=root_id).send()
        user_question_embedding = calculate_embedding([user_question.content])
        scores, indices = self.faiss_index.search(user_question_embedding, knowledge_top_k)
        # limits, distances, indices = self.faiss_index.range_search(user_question_embedding, threshold)
        question_knowledge_list = []
        related_question_list = [self.qa_knowledge['question_list'][i] for i in indices.tolist()[0]]
        # print(scores.tolist()[0])
        if scores.tolist()[0][0] > 120:
            related_question_list = []
        await cl.Message(author="Faiss", content=f'Searching result: {str(related_question_list)}', parent_id=root_id).send()

        await cl.Message(author="KG", content=f'Searching related knowledge.', parent_id=root_id).send()
        for knowledge in self.qa_knowledge['knowledge']:
            if knowledge[0] in related_question_list:
                question_knowledge_list.append(knowledge)
        for ontology in self.qa_knowledge['ontology']:
            if ontology[0] in related_question_list:
                question_knowledge_list.append(ontology)
        await cl.Message(author="KG", content=f'Searching result: {str(question_knowledge_list)}', parent_id=root_id).send()

        final_prompt = get_ticket_prompt(question_knowledge_list, user_question.content)
        await cl.Message(author="Prompt", content=f'{str(final_prompt)}', parent_id=root_id).send()
        llm_response = request_llm(final_prompt, response_history, **gen_kwargs)

        # print(len(llm_response['response'].replace('</s>', '')))
        await cl.Message(content=llm_response['response'].replace('</s>', '')).send()

        return llm_response['response'], llm_response['history']
    

    async def RAG_answer(self, user_question, response_history, **gen_kwargs):
        '''该if判断本次QA是否是RAG-based,若否,则跳转至simple_answer'''
        if user_question.elements==[]:
            return await self.simple_answer(user_question, response_history, **gen_kwargs)
        
        question=user_question.content
        file=user_question.elements[0]
        
        images_str=None
        prompt=None
        if file.name.split(".")[-1] in {"docx"}:
            file_stream = io.BytesIO(file.content)
            doc = docx.Document(file_stream)
            # Extracting text
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            full_text = '\n'.join(full_text)
            page_content=full_text

            '''下面这段被注释掉的代码是从docx文件中读取图片'''
            # # Extracting images
            # images = []
            # for rel in doc.part.rels.values():
            #     if "image" in rel.reltype:
            #         img = rel.target_part.blob
            #         images.append(img)   

            prompt=self.get_RAG_prompt(file,page_content,question)     
        elif file.name.split(".")[-1] in {"png","jpg"}:
            images_str=base64.b64encode(file.content).decode('utf-8')
            prompt=question
        elif file.name.split(".")[-1] in {"txt"} :                  
            page_content=file.content
            prompt=self.get_RAG_prompt(file,page_content,question)

        print(prompt)
        #LLM
        llm_response = request_llm(prompt=prompt, history=response_history, images_str=images_str, **gen_kwargs)

        await cl.Message(content=llm_response['response'].replace('</s>', '')).send()

        return llm_response['response'], llm_response['history']
    

    def get_RAG_prompt(self,file,page_content,question):
        metadata={"id":file.id, "name": file.name, "mime": file.mime}
        docs = [Document(page_content=page_content, metadata=metadata)]
        #document split
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
        splits = text_splitter.split_documents(docs)
        #document vectorstore
        embedding_function=SentenceTransformerEmbeddings(model_name=embedding_model_dict['m3e-base'])
        
        vectorstore = Chroma.from_documents(documents=splits, embedding=embedding_function)
        #document retrieve
        retriever = vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": 6})
        #prompt
        # prompt_template = hub.pull("rlm/rag-prompt")
        def prompt_template(context:str,question:str):
            return f"My question is: {question}\n\n There's some related information: {context}"

        def format_docs(docs):
            return "\n\n".join(doc.page_content for doc in docs)
        
        retrieved_docs = retriever.get_relevant_documents(question)
        context=format_docs(retrieved_docs)
        
        prompt=prompt_template(context, question)
        return prompt